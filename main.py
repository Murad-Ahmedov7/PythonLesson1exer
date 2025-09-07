from docx import Document
import re
doc=Document("data.docx")

result=doc.paragraphs[0].text

parts=re.split(r'(?<=[.!?])',result)



result_first=(parts[0].split())




result_first[0]=(result_first[0].capitalize())
result_first[3]=(result_first[3].lower())
result_first[4]=(result_first[4].replace("language","language,"))
result_first[7]=(result_first[7].lower())
result_first[9]=(result_first[9].replace("know,","know"))
result_first[13]=(result_first[13].replace("Correctly" ,  "correctly."))
result_first[14]=(result_first[14].capitalize())
result_first[20]=(result_first[20].replace("Happens." ,  "happen."))


final=''
for part in result_first:
    final+=' ' + part





result_second=(parts[1].split())

result_second[0]  = result_second[0].capitalize()
result_second[1]  = result_second[1].lower()
result_second[2]  = result_second[2].lower() + ","
result_second[4]  = result_second[4].lower()
result_second[7]  = result_second[7].lower()
result_second[-1] = result_second[-1].replace("checks?", "checks.")

final2=''
for part in result_second:
    final2+=' ' + part



# 3-cü cümlə
result_third = parts[3].split()

result_third[0]  = result_third[0].capitalize()
result_third[1]  = result_third[1].lower()
result_third[2]  = result_third[2].lower()
result_third[3]  = result_third[3].replace("developer", "developers")
result_third[4]  = result_third[4].lower()
result_third[6]  = result_third[6].lower()
result_third[11] = result_third[11].replace("maked;", "made,")
result_third[12] = result_third[12].replace("which", "which")
result_third[13] = result_third[13].replace("cause", "causes")
result_third[15] = result_third[15].lower()

final3=''
for part in result_third:
    final3+=' ' + part


the_last=final+final2+final3

print(the_last)


doc=Document()

doc.add_paragraph(the_last)

doc.save('correctData.docx')